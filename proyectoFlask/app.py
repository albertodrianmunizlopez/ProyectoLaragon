"""
Macuin Flask — Portal/Frontend.
Todas las rutas consumen datos de la API REST (FastAPI) vía HTTP.
"""
import os
import requests
from flask import Flask, render_template, request, redirect, url_for, session, abort, Response, jsonify

app = Flask(__name__)
app.secret_key = os.getenv("FLASK_SECRET_KEY", "macuin-flask-session-key-2026")

# ── URL base de la API (dentro de Docker: http://fastapi:8080) ──
API_URL = os.getenv("MACUIN_API_URL", "http://fastapi:8080")


# ============================================================
#  HELPERS — Llamadas a la API
# ============================================================

def api_get(endpoint, token=None, params=None):
    """GET a la API. Retorna JSON o None si falla."""
    headers = {}
    if token:
        headers["Authorization"] = f"Bearer {token}"
    try:
        r = requests.get(f"{API_URL}{endpoint}", headers=headers, params=params, timeout=5)
        if r.ok:
            return r.json()
    except requests.RequestException as e:
        print(f"[API ERROR] GET {endpoint}: {e}")
    return None


def api_post(endpoint, data=None, json_data=None, token=None):
    """POST a la API. Retorna (json, status_code)."""
    headers = {}
    if token:
        headers["Authorization"] = f"Bearer {token}"
    try:
        r = requests.post(f"{API_URL}{endpoint}", json=json_data, data=data, headers=headers, timeout=5)
        return r.json() if r.content else {}, r.status_code
    except requests.RequestException as e:
        print(f"[API ERROR] POST {endpoint}: {e}")
    return None, 500


def api_put(endpoint, json_data=None, token=None):
    """PUT a la API."""
    headers = {}
    if token:
        headers["Authorization"] = f"Bearer {token}"
    try:
        r = requests.put(f"{API_URL}{endpoint}", json=json_data, headers=headers, timeout=5)
        return r.json() if r.content else {}, r.status_code
    except requests.RequestException as e:
        print(f"[API ERROR] PUT {endpoint}: {e}")
    return None, 500


def api_patch(endpoint, json_data=None, token=None):
    """PATCH a la API."""
    headers = {}
    if token:
        headers["Authorization"] = f"Bearer {token}"
    try:
        r = requests.patch(f"{API_URL}{endpoint}", json=json_data, headers=headers, timeout=5)
        return r.json() if r.content else {}, r.status_code
    except requests.RequestException as e:
        print(f"[API ERROR] PATCH {endpoint}: {e}")
    return None, 500


def api_delete(endpoint, token=None):
    """DELETE a la API."""
    headers = {}
    if token:
        headers["Authorization"] = f"Bearer {token}"
    try:
        r = requests.delete(f"{API_URL}{endpoint}", headers=headers, timeout=5)
        return r.status_code
    except requests.RequestException as e:
        print(f"[API ERROR] DELETE {endpoint}: {e}")
    return 500


def get_token():
    """Obtener el token JWT de la sesión actual."""
    return session.get("token")


def get_admin_token():
    """Obtener el token JWT de la sesión admin."""
    return session.get("admin_token")


# ============================================================
#  LOGIN — Portal
# ============================================================

@app.route('/', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')

        # Llamar a la API de autenticación
        resp, status = api_post("/api/auth/login", json_data={
            "email": email,
            "password": password,
        })

        if status == 200 and resp and "access_token" in resp:
            session["token"] = resp["access_token"]
            session["usuario"] = resp.get("usuario", {})
            return redirect(url_for('inventario'))
        else:
            return render_template('login.html', error="Credenciales incorrectas")

    return render_template('login.html')


# ============================================================
#  INVENTARIO — Productos
# ============================================================

@app.route('/inventario')
def inventario():
    # Leer filtros y orden de la URL
    filtro_tipo = request.args.get('tipo', '')
    filtro_marca = request.args.get('marca', '')
    orden = request.args.get('orden', '')

    # Construir parámetros para la API
    params = {}
    if filtro_tipo:
        params['tipo'] = filtro_tipo
    if filtro_marca:
        params['marca'] = filtro_marca
    if orden:
        params['orden'] = orden

    # Obtener productos y catálogos de la API
    productos = api_get("/api/productos", params=params) or []
    tipos = api_get("/api/productos/catalogos/tipos") or []
    marcas = api_get("/api/productos/catalogos/marcas") or []

    # Adaptar formato para la plantilla
    autopartes = []
    for p in productos:
        autopartes.append({
            "codigo": p.get("codigo", ""),
            "nombre": p.get("nombre", ""),
            "imagen_url": p.get("imagen_url", ""),
            "tipo_nombre": p.get("tipo_nombre", ""),
            "marca_nombre": p.get("marca_nombre", ""),
            "stock": p.get("cantidad", 0),
            "precio": float(p.get("precio", 0)),
        })
    return render_template('inventario.html',
                           autopartes=autopartes,
                           tipos=tipos,
                           marcas=marcas,
                           filtro_tipo=filtro_tipo,
                           filtro_marca=filtro_marca,
                           orden=orden)


@app.route('/autoparte/nueva')
def nueva_autoparte():
    tipos = api_get("/api/productos/catalogos/tipos") or []
    marcas = api_get("/api/productos/catalogos/marcas") or []
    return render_template('formulario_autoparte.html', autoparte=None, tipos=tipos, marcas=marcas)


@app.route('/autoparte/editar/<codigo>')
def editar_autoparte(codigo):
    producto = api_get(f"/api/productos/buscar", params={"codigo": codigo})
    tipos = api_get("/api/productos/catalogos/tipos") or []
    marcas = api_get("/api/productos/catalogos/marcas") or []

    if producto:
        autoparte = {
            "codigo": producto.get("codigo", codigo),
            "nombre": producto.get("nombre", ""),
            "imagen_url": producto.get("imagen_url", ""),
            "id_tipo_autoparte": producto.get("id_tipo_autoparte"),
            "id_marca": producto.get("id_marca"),
            "stock": producto.get("cantidad", 0),
            "precio": float(producto.get("precio", 0)),
            "descripcion": producto.get("descripcion", ""),
        }
    else:
        autoparte = {"codigo": codigo, "nombre": "", "imagen_url": "", "id_tipo_autoparte": None, "id_marca": None, "stock": 0, "precio": 0, "descripcion": ""}

    return render_template('formulario_autoparte.html', autoparte=autoparte, tipos=tipos, marcas=marcas)


@app.route('/autoparte/guardar', methods=['POST'])
def guardar_autoparte():
    token = get_token() or get_admin_token()
    datos_form = request.form

    codigo = datos_form.get('codigo', '').strip()

    # Manejar tipo: si es "nuevo", crear primero vía API
    id_tipo = datos_form.get('id_tipo_autoparte', '')
    if id_tipo == 'nuevo':
        nuevo_nombre = datos_form.get('nuevo_tipo', '').strip()
        if nuevo_nombre:
            resp, st = api_post("/api/productos/catalogos/tipos", json_data={"nombre": nuevo_nombre}, token=token)
            if resp and resp.get("id"):
                id_tipo = resp["id"]
            else:
                id_tipo = None
        else:
            id_tipo = None
    else:
        id_tipo = int(id_tipo) if id_tipo else None

    # Manejar marca: si es "nuevo", crear primero vía API
    id_marca = datos_form.get('id_marca', '')
    if id_marca == 'nuevo':
        nuevo_nombre = datos_form.get('nueva_marca', '').strip()
        if nuevo_nombre:
            resp, st = api_post("/api/productos/catalogos/marcas", json_data={"nombre": nuevo_nombre}, token=token)
            if resp and resp.get("id"):
                id_marca = resp["id"]
            else:
                id_marca = None
        else:
            id_marca = None
    else:
        id_marca = int(id_marca) if id_marca else None

    json_data = {
        "nombre": datos_form.get('nombre', ''),
        "descripcion": datos_form.get('descripcion', ''),
        "imagen_url": datos_form.get('imagen_url', '').strip() or None,
        "id_tipo_autoparte": id_tipo,
        "id_marca": id_marca,
        "cantidad": int(datos_form.get('stock', 0)),
        "estatus_producto": "en_stock",
        "precio": float(datos_form.get('precio', 0)),
    }

    if codigo:
        # Editando un producto existente
        existente = api_get(f"/api/productos/buscar", params={"codigo": codigo})
        if existente and existente.get("id"):
            json_data["codigo"] = codigo
            api_put(f"/api/productos/{existente['id']}", json_data=json_data, token=token)
    else:
        # Nuevo producto — el código lo genera la API automáticamente
        api_post("/api/productos", json_data=json_data, token=token)

    return redirect(url_for('inventario'))


@app.route('/autoparte/eliminar/<codigo>', methods=['POST'])
def eliminar_autoparte(codigo):
    token = get_token() or get_admin_token()
    # Buscar producto por código
    producto = api_get(f"/api/productos/buscar", params={"codigo": codigo})
    if producto and producto.get("id"):
        api_delete(f"/api/productos/{producto['id']}", token=token)
    return redirect(url_for('inventario'))


# ============================================================
#  PEDIDOS
# ============================================================

@app.route('/pedidos')
def dashboard_pedidos():
    token = get_token() or get_admin_token()

    # Leer filtros de la URL
    filtro_estatus = request.args.get('estatus', '')
    orden = request.args.get('orden', '')
    fecha_inicio = request.args.get('fecha_inicio', '')
    fecha_fin = request.args.get('fecha_fin', '')
    filtro_mes = request.args.get('mes', '')
    filtro_anio = request.args.get('anio', '')

    # Construir parámetros para la API
    params = {}
    if filtro_estatus:
        params['estado'] = filtro_estatus
    if orden:
        params['orden'] = orden
    if fecha_inicio:
        params['fecha_inicio'] = fecha_inicio
    if fecha_fin:
        params['fecha_fin'] = fecha_fin
    if filtro_mes:
        params['mes'] = filtro_mes
    if filtro_anio:
        params['anio'] = filtro_anio

    resp = api_get("/api/pedidos", params=params, token=token) or {"pedidos": [], "total": 0}

    pedidos = []
    for p in resp.get("pedidos", []):
        fecha_raw = p.get("fecha_pedido", "")
        fecha = fecha_raw[:10] if fecha_raw else ""
        pedidos.append({
            "id": p.get("codigo_pedido", ""),
            "fecha": fecha,
            "cliente": {
                "nombre": p.get("usuario_nombre", f"Usuario #{p.get('id_usuario', '?')}"),
            },
            "articulos": p.get("num_articulos", 0),
            "total": float(p.get("total", 0)),
            "estatus": p.get("estado_pedido", "pendiente").capitalize(),
        })

    total_pedidos = resp.get("total", len(pedidos))

    return render_template('pedidos_dashboard.html',
                           pedidos=pedidos,
                           total_pedidos=total_pedidos,
                           filtro_estatus=filtro_estatus,
                           orden=orden,
                           fecha_inicio=fecha_inicio,
                           fecha_fin=fecha_fin,
                           filtro_mes=filtro_mes,
                           filtro_anio=filtro_anio)


@app.route('/pedidos/detalle/<id>')
def detalle_pedido(id):
    token = get_token() or get_admin_token()

    # El template usa el código de pedido como ID, necesitamos buscar por ID numérico
    # Intentar parsear como número; si falla, buscar en la lista
    pedido_id = id
    try:
        pedido_id = int(id)
    except ValueError:
        # Es un código como ORD-9011, buscar en la lista
        resp = api_get("/api/pedidos", token=token) or {"pedidos": []}
        for p in resp.get("pedidos", []):
            if p.get("codigo_pedido") == id:
                pedido_id = p.get("id")
                break

    detalle = api_get(f"/api/pedidos/{pedido_id}", token=token)

    if not detalle:
        abort(404)

    # Adaptar al formato que espera la plantilla
    pedido = {
        "id": detalle.get("codigo_pedido", id),
        "fecha_creacion": detalle.get("fecha_pedido", ""),
        "estatus": detalle.get("estado_pedido", "pendiente"),
        "cliente": {
            "nombre": detalle.get("usuario_nombre", ""),
            "telefono": "",
            "email": detalle.get("usuario_email", ""),
            "direccion": detalle.get("direccion_completa", ""),
        },
        "articulos": [],
        "finanzas": {
            "subtotal": float(detalle.get("subtotal", 0)),
            "impuestos": float(detalle.get("impuestos", 0)),
            "total": float(detalle.get("total", 0)),
        },
    }

    for art in detalle.get("productos", []):
        pedido["articulos"].append({
            "codigo": art.get("producto_codigo", ""),
            "nombre": art.get("producto_nombre", ""),
            "cantidad": art.get("cantidad", 0),
            "precio_unitario": float(art.get("precio_unitario", 0)),
        })

    return render_template('detalle_pedido.html', pedido=pedido)


@app.route('/pedidos/actualizar', methods=['POST'])
def actualizar_estatus_pedido():
    token = get_token() or get_admin_token()
    pedido_id = request.form.get('pedido_id')
    nuevo_estatus = request.form.get('estatus')

    # Buscar ID numérico si nos pasan código
    pid = pedido_id
    try:
        pid = int(pedido_id)
    except ValueError:
        resp = api_get("/api/pedidos", token=token) or {"pedidos": []}
        for p in resp.get("pedidos", []):
            if p.get("codigo_pedido") == pedido_id:
                pid = p.get("id")
                break

    api_patch(f"/api/pedidos/{pid}/estado", json_data={"estado": nuevo_estatus}, token=token)

    return redirect(url_for('detalle_pedido', id=pedido_id))


# ============================================================
#  REPORTES
# ============================================================

@app.route('/reportes')
def vista_reportes():
    token = get_token() or get_admin_token()

    kpis = api_get("/api/reportes/kpis", token=token) or {
        "ventas": 0, "pedidos": 0, "clientes": 0,
        "pendientes": 0, "entregados": 0, "cancelados": 0,
        "productos": 0, "stock_total": 0,
    }

    return render_template('reportes.html', kpi=kpis)


@app.route('/reportes/crear')
def crear_reporte():
    token = get_token() or get_admin_token()
    tipo = request.args.get('tipo', '')
    fecha_inicio = request.args.get('fecha_inicio', '')
    fecha_fin = request.args.get('fecha_fin', '')
    estado = request.args.get('estado', '')
    marca_id = request.args.get('marca_id', '')
    top = request.args.get('top', '')
    tipo_id = request.args.get('tipo_id', '')
    ordenar = request.args.get('ordenar', '')
    precio_min = request.args.get('precio_min', '')
    precio_max = request.args.get('precio_max', '')
    alfa = request.args.get('alfa', '')
    generado = request.args.get('generado', '')

    # Cargar marcas y tipos para los dropdowns
    marcas = api_get("/api/productos/catalogos/marcas", token=token) or []
    tipos = api_get("/api/productos/catalogos/tipos", token=token) or []

    datos = []
    total_datos = 0

    if generado and tipo:
        params = {"tipo": tipo}
        if fecha_inicio:
            params['fecha_inicio'] = fecha_inicio
        if fecha_fin:
            params['fecha_fin'] = fecha_fin
        if estado:
            params['estado'] = estado
        if marca_id:
            params['marca_id'] = marca_id
        if top:
            params['top'] = top
        if tipo_id:
            params['tipo_id'] = tipo_id
        if ordenar:
            params['ordenar'] = ordenar
        if precio_min:
            params['precio_min'] = precio_min
        if precio_max:
            params['precio_max'] = precio_max
        if alfa:
            params['alfa'] = alfa

        resp = api_get("/api/reportes/datos", params=params, token=token) or {"datos": [], "total": 0}
        datos = resp.get("datos", [])
        total_datos = resp.get("total", 0)

    return render_template('crear_reporte.html',
                           tipo=tipo,
                           fecha_inicio=fecha_inicio,
                           fecha_fin=fecha_fin,
                           estado=estado,
                           marca_id=marca_id,
                           top=top,
                           tipo_id=tipo_id,
                           ordenar=ordenar,
                           precio_min=precio_min,
                           precio_max=precio_max,
                           alfa=alfa,
                           generado=generado,
                           datos=datos,
                           total_datos=total_datos,
                           marcas=marcas,
                           tipos=tipos)


@app.route('/reportes/exportar')
def exportar_reporte():
    """Generar y descargar reporte en PDF, XLSX o DOCX."""
    token = get_token() or get_admin_token()
    formato = request.args.get('formato', 'xlsx')
    tipo = request.args.get('tipo', 'ventas')
    fecha_inicio = request.args.get('fecha_inicio', '')
    fecha_fin = request.args.get('fecha_fin', '')
    estado = request.args.get('estado', '')
    marca_id = request.args.get('marca_id', '')
    top = request.args.get('top', '')
    tipo_id = request.args.get('tipo_id', '')
    ordenar = request.args.get('ordenar', '')
    precio_min = request.args.get('precio_min', '')
    precio_max = request.args.get('precio_max', '')
    alfa = request.args.get('alfa', '')

    params = {"tipo": tipo}
    if fecha_inicio:
        params['fecha_inicio'] = fecha_inicio
    if fecha_fin:
        params['fecha_fin'] = fecha_fin
    if estado:
        params['estado'] = estado
    if marca_id:
        params['marca_id'] = marca_id
    if top:
        params['top'] = top
    if tipo_id:
        params['tipo_id'] = tipo_id
    if ordenar:
        params['ordenar'] = ordenar
    if precio_min:
        params['precio_min'] = precio_min
    if precio_max:
        params['precio_max'] = precio_max
    if alfa:
        params['alfa'] = alfa

    resp = api_get("/api/reportes/datos", params=params, token=token) or {"datos": []}
    datos = resp.get("datos", [])

    titulo_map = {
        "ventas": "Reporte de Ventas",
        "pedidos": "Reporte de Pedidos",
        "inventario": "Reporte de Inventario",
        "clientes": "Reporte de Clientes",
    }
    titulo = titulo_map.get(tipo, "Reporte")

    periodo = ""
    if fecha_inicio and fecha_fin:
        periodo = f"Periodo: {fecha_inicio} a {fecha_fin}"
    elif fecha_inicio:
        periodo = f"Desde: {fecha_inicio}"
    elif fecha_fin:
        periodo = f"Hasta: {fecha_fin}"

    if formato == 'xlsx':
        return _exportar_xlsx(datos, tipo, titulo, periodo)
    elif formato == 'pdf':
        return _exportar_pdf(datos, tipo, titulo, periodo)
    elif formato == 'docx':
        return _exportar_docx(datos, tipo, titulo, periodo)
    else:
        return "Formato no soportado", 400


def _get_headers_and_rows(datos, tipo):
    """Obtener headers y filas segun tipo de reporte."""
    if tipo == "ventas":
        headers = ["Codigo", "Fecha", "Cliente", "Articulos", "Subtotal", "Impuestos", "Total", "Estado"]
        rows = [[d.get("codigo",""), d.get("fecha",""), d.get("cliente",""),
                 str(d.get("articulos",0)), f"${d.get('subtotal',0):,.2f}",
                 f"${d.get('impuestos',0):,.2f}", f"${d.get('total',0):,.2f}",
                 d.get("estado","")] for d in datos]
    elif tipo == "pedidos":
        headers = ["Codigo", "Fecha", "Cliente", "Articulos", "Total", "Estado"]
        rows = [[d.get("codigo",""), d.get("fecha",""), d.get("cliente",""),
                 str(d.get("articulos",0)), f"${d.get('total',0):,.2f}",
                 d.get("estado","")] for d in datos]
    elif tipo == "inventario":
        headers = ["Codigo", "Nombre", "Tipo", "Marca", "Precio", "Stock"]
        rows = [[d.get("codigo",""), d.get("nombre",""), d.get("tipo",""),
                 d.get("marca",""), f"${d.get('precio',0):,.2f}",
                 str(d.get("stock",0))] for d in datos]
    elif tipo == "clientes":
        headers = ["Nombre", "Email", "Pedidos", "Gasto Total"]
        rows = [[d.get("nombre",""), d.get("email",""),
                 str(d.get("pedidos",0)), f"${d.get('gasto_total',0):,.2f}"] for d in datos]
    else:
        headers = []
        rows = []
    return headers, rows


def _exportar_xlsx(datos, tipo, titulo, periodo):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    import io
    from datetime import datetime

    # Macuin brand colors
    DARK = "1A1A2E"
    RED = "E9302A"
    DARK_ROW = "0E0E1A"
    ALT_ROW = "16162B"
    LIGHTER = "2A2A40"
    WHITE = "F0E6EA"
    MUTED = "A09298"

    wb = Workbook()
    ws = wb.active
    ws.title = titulo[:31]
    ws.sheet_properties.tabColor = RED

    headers, rows = _get_headers_and_rows(datos, tipo)
    num_cols = len(headers) if headers else 1

    # ── Brand header ──
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=num_cols)
    ws['A1'] = "MACUIN AUTOPARTES"
    ws['A1'].font = Font(bold=True, size=16, color=WHITE)
    ws['A1'].fill = PatternFill(start_color=DARK, end_color=DARK, fill_type="solid")
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center')

    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=num_cols)
    ws['A2'] = titulo.upper()
    ws['A2'].font = Font(bold=True, size=12, color=RED)
    ws['A2'].fill = PatternFill(start_color=DARK, end_color=DARK, fill_type="solid")
    ws['A2'].alignment = Alignment(horizontal='center')

    info_row = 3
    if periodo:
        ws.merge_cells(start_row=info_row, start_column=1, end_row=info_row, end_column=num_cols)
        ws.cell(row=info_row, column=1, value=periodo).font = Font(italic=True, size=10, color=MUTED)
        ws.cell(row=info_row, column=1).fill = PatternFill(start_color=DARK, end_color=DARK, fill_type="solid")
        ws.cell(row=info_row, column=1).alignment = Alignment(horizontal='center')
        info_row += 1

    ws.merge_cells(start_row=info_row, start_column=1, end_row=info_row, end_column=num_cols)
    ws.cell(row=info_row, column=1, value=f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}").font = Font(size=9, color=MUTED)
    ws.cell(row=info_row, column=1).fill = PatternFill(start_color=DARK, end_color=DARK, fill_type="solid")
    ws.cell(row=info_row, column=1).alignment = Alignment(horizontal='center')

    # ── Data headers ──
    start_row = info_row + 2
    header_font = Font(bold=True, color=WHITE, size=10)
    header_fill = PatternFill(start_color=RED, end_color=RED, fill_type="solid")
    border_bottom = Border(bottom=Side(style='thin', color=LIGHTER))

    for col_idx, h in enumerate(headers, 1):
        cell = ws.cell(row=start_row, column=col_idx, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center', vertical='center')

    # ── Data rows with alternating colors ──
    row_fill_a = PatternFill(start_color=DARK_ROW, end_color=DARK_ROW, fill_type="solid")
    row_fill_b = PatternFill(start_color=ALT_ROW, end_color=ALT_ROW, fill_type="solid")
    data_font = Font(size=10, color=WHITE)

    for row_idx, row_data in enumerate(rows, start_row + 1):
        fill = row_fill_a if (row_idx % 2 == 0) else row_fill_b
        for col_idx, val in enumerate(row_data, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=str(val))
            cell.font = data_font
            cell.fill = fill
            cell.border = border_bottom
            cell.alignment = Alignment(horizontal='center')

    # ── Footer ──
    footer_row = start_row + len(rows) + 2
    ws.merge_cells(start_row=footer_row, start_column=1, end_row=footer_row, end_column=num_cols)
    ws.cell(row=footer_row, column=1, value=f"Total de registros: {len(rows)}").font = Font(bold=True, size=10, color=RED)

    # Auto-width
    for col_idx in range(1, num_cols + 1):
        max_len = 0
        for r in range(start_row, start_row + 1 + len(rows)):
            cell = ws.cell(row=r, column=col_idx)
            if cell.value:
                max_len = max(max_len, len(str(cell.value)))
        ws.column_dimensions[get_column_letter(col_idx)].width = min(max_len + 4, 45)

    # Row heights
    ws.row_dimensions[1].height = 28
    ws.row_dimensions[2].height = 22

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)

    filename = f"macuin_{tipo}.xlsx"
    response = Response(
        output.getvalue(),
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    )
    response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def _exportar_pdf(datos, tipo, titulo, periodo):
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.lib.units import cm, mm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT
    import io
    from datetime import datetime

    # Macuin brand colors — dark theme
    BG_PAGE = colors.HexColor('#0A0A14')
    BRAND_DARK = colors.HexColor('#1A1A2E')
    BRAND_RED = colors.HexColor('#E9302A')
    ROW_A = colors.HexColor('#0E0E1A')
    ROW_B = colors.HexColor('#16162B')
    BRAND_LIGHT = colors.HexColor('#F0E6EA')
    BRAND_MUTED = colors.HexColor('#A09298')
    GRID_COLOR = colors.HexColor('#2A2A40')

    output = io.BytesIO()
    page_w, page_h = landscape(A4)

    def draw_bg(canvas, doc):
        canvas.saveState()
        canvas.setFillColor(BG_PAGE)
        canvas.rect(0, 0, page_w, page_h, fill=1, stroke=0)
        canvas.restoreState()

    doc = SimpleDocTemplate(output, pagesize=landscape(A4),
                            leftMargin=1.5*cm, rightMargin=1.5*cm,
                            topMargin=1.5*cm, bottomMargin=1.5*cm)
    elements = []
    styles = getSampleStyleSheet()
    tbl_width = page_w - 3*cm

    # Title bar
    title_data = [["MACUIN AUTOPARTES"]]
    title_tbl = Table(title_data, colWidths=[tbl_width])
    title_tbl.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), BRAND_DARK),
        ('TEXTCOLOR', (0, 0), (-1, -1), BRAND_LIGHT),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 16),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('TOPPADDING', (0, 0), (-1, -1), 12),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
    ]))
    elements.append(title_tbl)

    # Subtitle bar
    sub_text = titulo.upper()
    if periodo:
        sub_text += f"  |  {periodo}"
    sub_data = [[sub_text]]
    sub_tbl = Table(sub_data, colWidths=[tbl_width])
    sub_tbl.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), BRAND_RED),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.white),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    elements.append(sub_tbl)
    elements.append(Spacer(1, 0.5*cm))

    headers, rows = _get_headers_and_rows(datos, tipo)

    if rows:
        safe_rows = [[str(v) for v in row] for row in rows]
        table_data = [headers] + safe_rows
        table = Table(table_data, repeatRows=1)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), BRAND_DARK),
            ('TEXTCOLOR', (0, 0), (-1, 0), BRAND_LIGHT),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 9),
            ('TEXTCOLOR', (0, 1), (-1, -1), BRAND_LIGHT),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 8),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
            ('TOPPADDING', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
            ('TOPPADDING', (0, 1), (-1, -1), 6),
            ('LINEBELOW', (0, 0), (-1, 0), 2, BRAND_RED),
            ('GRID', (0, 1), (-1, -1), 0.25, GRID_COLOR),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [ROW_A, ROW_B]),
        ]))
        elements.append(table)
    else:
        no_data_style = ParagraphStyle('NoData', parent=styles['Normal'], alignment=TA_CENTER, fontSize=11, textColor=BRAND_LIGHT)
        elements.append(Paragraph("Sin datos para el periodo seleccionado.", no_data_style))

    elements.append(Spacer(1, 0.8*cm))

    # Footer bar
    footer_text = f"Total: {len(rows)} registros  |  Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}  |  MACUIN Autopartes"
    footer_data = [[footer_text]]
    footer_tbl = Table(footer_data, colWidths=[tbl_width])
    footer_tbl.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), BRAND_DARK),
        ('TEXTCOLOR', (0, 0), (-1, -1), BRAND_MUTED),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
    ]))
    elements.append(footer_tbl)

    doc.build(elements, onFirstPage=draw_bg, onLaterPages=draw_bg)
    output.seek(0)

    filename = f"macuin_{tipo}.pdf"
    response = Response(
        output.getvalue(),
        mimetype='application/pdf',
    )
    response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def _exportar_docx(datos, tipo, titulo, periodo):
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io
    from datetime import datetime

    BRAND_DARK = RGBColor(26, 26, 46)
    BRAND_RED = RGBColor(233, 48, 42)
    BRAND_WHITE = RGBColor(240, 230, 234)

    def set_cell_bg(cell, color_hex):
        """Set background color of a docx cell."""
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color_hex)
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)

    doc = DocxDocument()

    # Set narrow margins
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(1.5)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("MACUIN AUTOPARTES")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = BRAND_DARK

    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run(titulo.upper())
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = BRAND_RED

    if periodo:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(periodo)
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(120, 120, 130)

    # Timestamp
    ts_p = doc.add_paragraph()
    ts_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ts_p.add_run(f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(150, 150, 160)

    doc.add_paragraph("")

    headers, rows = _get_headers_and_rows(datos, tipo)

    if rows:
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            set_cell_bg(cell, "1A1A2E")
            run = cell.paragraphs[0].add_run(h)
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = BRAND_WHITE
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Data rows with alternating dark shade
        for row_idx, row_data in enumerate(rows):
            row_cells = table.add_row().cells
            bg = "16162B" if row_idx % 2 == 1 else "0E0E1A"
            for i, val in enumerate(row_data):
                row_cells[i].text = ""
                set_cell_bg(row_cells[i], bg)
                run = row_cells[i].paragraphs[0].add_run(str(val))
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(220, 210, 215)
                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Sin datos para el periodo seleccionado.")
        run.font.size = Pt(11)

    doc.add_paragraph("")

    # Footer
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run(f"Total: {len(rows)} registros  |  MACUIN Autopartes")
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = BRAND_RED

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    filename = f"macuin_{tipo}.docx"
    response = Response(
        output.getvalue(),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


# ============================================================
#  PORTAL DE ADMINISTRACIÓN — Gestión de Usuarios
# ============================================================

@app.route('/admin/login', methods=['GET', 'POST'])
def admin_login():
    mensaje_error = None
    if request.method == 'POST':
        email = request.form.get('email', '')
        password = request.form.get('password', '')

        # Autenticarse contra la API
        resp, status = api_post("/api/auth/login", json_data={
            "email": email,
            "password": password,
        })

        if status == 200 and resp and "access_token" in resp:
            usuario = resp.get("usuario", {})
            # Verificar que sea admin o superadmin
            if usuario.get("status") in ("administrador", "superadministrador"):
                session["admin_token"] = resp["access_token"]
                session["admin_usuario"] = usuario
                return redirect(url_for('admin_home'))
            else:
                mensaje_error = 'No tienes permisos de administrador.'
        else:
            mensaje_error = 'Credenciales incorrectas. Verifica tu correo y contraseña.'

    return render_template('admin_login.html', error=mensaje_error)


@app.route('/admin')
def admin_home():
    token = get_admin_token()
    if not token:
        return redirect(url_for('admin_login'))

    # Obtener stats reales de la API
    stats_data = api_get("/api/usuarios/stats", token=token)
    stats = stats_data if stats_data else {'total': 0, 'admins': 0, 'usuarios': 0, 'activos': 0}

    # Mapear nombres para compatibilidad con la plantilla
    stats['clientes'] = stats.get('usuarios', 0)

    return render_template('admin_home.html', stats=stats)


@app.route('/admin/usuarios')
def admin_usuarios():
    token = get_admin_token()
    if not token:
        return redirect(url_for('admin_login'))

    filtro = request.args.get('filtro', 'todos')

    # Mapear filtros del template a los de la API
    filtro_api = None
    if filtro == 'superadmin':
        filtro_api = 'superadministrador'
    elif filtro == 'admin':
        filtro_api = 'administrador'
    elif filtro == 'cliente':
        filtro_api = 'usuario'

    # Obtener usuarios de la API
    usuarios_raw = api_get("/api/usuarios", token=token, params={"filtro": filtro_api} if filtro_api else None) or []

    # Adaptar formato para la plantilla
    usuarios = []
    for u in usuarios_raw:
        status_raw = u.get("status", "usuario")
        if status_raw == "superadministrador":
            rol = "superadmin"
        elif status_raw == "administrador":
            rol = "admin"
        else:
            rol = "cliente"
        usuarios.append({
            "id": u.get("id"),
            "nombre": f"{u.get('nombre', '')} {u.get('apellidos', '')}",
            "email": u.get("email", ""),
            "rol": rol,
            "fecha_alta": u.get("created_at", "")[:10] if u.get("created_at") else "",
            "activo": u.get("activo", False),
        })

    # Obtener totales para las tabs
    stats = api_get("/api/usuarios/stats", token=token) or {}
    totales = {
        'todos': stats.get('total', len(usuarios)),
        'superadmin': stats.get('superadmins', 0),
        'admin': stats.get('admins', 0),
        'cliente': stats.get('usuarios', 0),
    }

    mensaje = request.args.get('msg', '')
    return render_template('admin_usuarios.html', usuarios=usuarios, filtro=filtro, totales=totales, mensaje=mensaje)


@app.route('/admin/usuarios/editar/<int:usuario_id>', methods=['GET', 'POST'])
def admin_editar_usuario(usuario_id):
    token = get_admin_token()
    if not token:
        return redirect(url_for('admin_login'))

    if request.method == 'POST':
        nuevo_nombre = request.form.get('nombre', '').strip()
        nuevo_email = request.form.get('email', '').strip()
        nuevo_rol = request.form.get('rol', '')
        nuevo_activo = request.form.get('activo') == 'on'
        nueva_password = request.form.get('password', '').strip()

        # Separar nombre y apellidos si vienen juntos
        partes_nombre = nuevo_nombre.split(' ', 1) if nuevo_nombre else ['', '']
        nombre = partes_nombre[0]
        apellidos = partes_nombre[1] if len(partes_nombre) > 1 else ''

        # Mapear rol del template al estatus de la API
        rol_map = {'superadmin': 'superadministrador', 'admin': 'administrador', 'cliente': 'usuario'}
        status_api = rol_map.get(nuevo_rol, 'usuario')

        json_data = {
            "nombre": nombre,
            "apellidos": apellidos,
            "email": nuevo_email,
            "telefono": request.form.get('telefono', '').strip() or None,
            "status": status_api,
            "activo": nuevo_activo,
        }
        if nueva_password:
            json_data["password"] = nueva_password

        api_put(f"/api/usuarios/{usuario_id}", json_data=json_data, token=token)
        return redirect(url_for('admin_usuarios', msg='editado'))

    # GET → obtener usuario de la API (endpoint admin incluye password_plain)
    u = api_get(f"/api/usuarios/{usuario_id}/admin", token=token)
    if not u:
        # Fallback al endpoint normal si el admin no existe
        u = api_get(f"/api/usuarios/{usuario_id}", token=token)
    if not u:
        abort(404)

    status_raw = u.get("status", "usuario")
    usuario = {
        "id": u.get("id"),
        "nombre": f"{u.get('nombre', '')} {u.get('apellidos', '')}",
        "email": u.get("email", ""),
        "telefono": u.get("telefono", "") or "",
        "rol": "admin" if status_raw in ("administrador", "superadministrador") else "cliente",
        "status_detalle": status_raw,
        "activo": u.get("activo", False),
        "password_plain": u.get("password_plain", ""),
        "fecha_alta": u.get("created_at", "")[:10] if u.get("created_at") else "No disponible",
        "updated_at": u.get("updated_at", "")[:10] if u.get("updated_at") else "",
    }
    return render_template('editar_usuario.html', usuario=usuario)


@app.route('/admin/usuarios/eliminar/<int:usuario_id>', methods=['POST'])
def admin_eliminar_usuario(usuario_id):
    token = get_admin_token()
    api_delete(f"/api/usuarios/{usuario_id}", token=token)
    return redirect(url_for('admin_usuarios', msg='eliminado'))


@app.route('/admin/usuarios/cambiar-rol', methods=['POST'])
def admin_cambiar_rol():
    token = get_admin_token()
    usuario_id = int(request.form.get('usuario_id', 0))
    nuevo_rol = request.form.get('nuevo_rol')

    # Mapear al estatus de la API
    rol_map = {'superadmin': 'superadministrador', 'admin': 'administrador', 'cliente': 'usuario'}
    status_api = rol_map.get(nuevo_rol, 'usuario')
    api_patch(f"/api/usuarios/{usuario_id}/cambiar-rol", json_data={"nuevo_rol": status_api}, token=token)
    return redirect(url_for('admin_usuarios'))


@app.route('/admin/usuarios/toggle-estado', methods=['POST'])
def admin_toggle_estado():
    token = get_admin_token()
    usuario_id = request.form.get('usuario_id')
    api_patch(f"/api/usuarios/{usuario_id}/toggle-estado", token=token)
    return redirect(url_for('admin_usuarios'))


@app.route('/admin/usuarios/nuevo', methods=['POST'])
def admin_nuevo_usuario():
    token = get_admin_token()
    datos = request.form

    nombre = datos.get('nombre', '').strip()
    apellidos = datos.get('apellidos', '').strip()
    rol = datos.get('rol', 'cliente')
    # Mapear rol del formulario al status de la API
    rol_map = {'superadmin': 'superadministrador', 'admin': 'administrador', 'cliente': 'usuario'}
    status_api = rol_map.get(rol, 'usuario')

    json_data = {
        "nombre": nombre,
        "apellidos": apellidos,
        "email": datos.get('email', ''),
        "password": datos.get('password', 'macuin2026'),
        "telefono": datos.get('telefono', '').strip() or None,
        "status": status_api,
    }
    resp, status_code = api_post("/api/usuarios", json_data=json_data, token=token)

    if status_code == 409:
        return redirect(url_for('admin_usuarios', msg='email_duplicado'))

    return redirect(url_for('admin_usuarios', msg='creado'))


@app.route('/admin/api/check-email')
def admin_check_email():
    """Proxy para verificar email duplicado vía la API."""
    email = request.args.get('email', '')
    token = get_admin_token()
    result = api_get(f"/api/auth/check-email?email={email}", token=token)
    if result:
        return jsonify(result)
    return jsonify({"exists": False})


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
